import os
import re
import json
import base64
import requests
import tempfile
import webbrowser
import pyautogui
import speech_recognition as sr
from pathlib import Path
from datetime import datetime
from PIL import Image
from groq import Groq
from voice import speak
from nlp_router import classify, detect_subject_nlp

# ── Central config (API keys live here) ──
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import GROQ_API_KEY as _CONFIG_GROQ_KEY

# ── ai_router: smart Ollama/Groq routing ──
from ai_router import ask_ai as _routed_ask_ai

# ================= ⚙️ CONFIG =================

GRADE        = 11
BOARD        = "CBSE"
SUBJECTS     = ["maths", "physics", "chemistry"]
WATCH_FOLDER = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR   = os.path.join(tempfile.gettempdir(), "FRIDAY_STUDY")

os.makedirs(OUTPUT_DIR, exist_ok=True)

print(f"[CONFIG] Grade: {GRADE} | Board: {BOARD}")
print(f"[CONFIG] Output directory: {OUTPUT_DIR}")

# ================= 🔑 API SETUP =================

# Groq — used for image vision, complex study tasks
# Key imported from config.py so there is one place to update it
GROQ_API_KEY        = _CONFIG_GROQ_KEY or os.getenv("GROQ_API_KEY", "")
GROQ_TEXT_MODEL     = "llama-3.3-70b-versatile"   # best Groq text model
GROQ_VISION_MODEL   = "meta-llama/llama-4-scout-17b-16e-instruct"

try:
    groq_client = Groq(api_key=GROQ_API_KEY)
    print(f"[GROQ] Client ready | text: {GROQ_TEXT_MODEL} | vision: {GROQ_VISION_MODEL}")
except Exception as e:
    groq_client = None
    print(f"[GROQ] Failed to initialize: {e}")

# Graphs and simulations now use Groq llama-3.3-70b (same client, no extra setup needed)

# ================= 🧠 SESSION STATE =================

SESSION = {
    "last_topic":        None,
    "last_subject":      None,
    "last_solution":     None,
    "last_quiz_file":    None,
    "last_wrong_topics": [],
    "last_image":        None,
}

print("[SESSION] Session state initialized.")

# ================= 🧠 SUBJECT DETECTION =================

SUBJECT_KEYWORDS = {
    "maths":     ["math", "maths", "algebra", "calculus", "trigonometry",
                  "geometry", "equation", "integral", "derivative", "matrix",
                  "vector", "probability", "statistics", "logarithm", "binomial",
                  "differentiate", "integrate", "limit", "function"],
    "physics":   ["physics", "force", "motion", "velocity", "acceleration",
                  "energy", "momentum", "wave", "optics", "electricity",
                  "magnetism", "current", "resistance", "circuit", "gravity",
                  "thermodynamics", "nuclear", "quantum", "projectile",
                  "pendulum", "oscillation", "refraction", "lens"],
    "chemistry": ["chemistry", "reaction", "molecule", "atom", "bond",
                  "acid", "base", "organic", "inorganic", "periodic",
                  "element", "compound", "mole", "titration", "equilibrium",
                  "electrochemistry", "polymer", "hydrocarbon", "alkane",
                  "aldehyde", "ketone", "ester", "salt"],
}

def detect_subject(command):
    """NLP-based subject detection with keyword fallback."""
    # try NLP first
    subj = detect_subject_nlp(command)
    if subj:
        print(f"[SUBJECT] NLP detected: '{subj}'")
        SESSION["last_subject"] = subj
        return subj

    # fallback to keyword matching for short commands
    cmd = command.lower()
    for subject, keywords in SUBJECT_KEYWORDS.items():
        for kw in keywords:
            if kw in cmd:
                print(f"[SUBJECT] Keyword detected: '{subject}' via '{kw}'")
                SESSION["last_subject"] = subject
                return subject

    fallback = SESSION.get("last_subject") or "general science"
    print(f"[SUBJECT] No match — fallback: '{fallback}'")
    return fallback

# ================= 🔍 CHAIN DETECTION =================

def detect_chains(command):
    cmd = command.lower()
    chains = []

    screen_words = ["on my screen", "on screen", "my screen", "on the screen",
                    "from my screen", "in my screen"]
    wants_screen   = any(w in cmd for w in screen_words)
    # exclude homework commands from solve chain — they go to solve_homework() instead
    is_homework_cmd = any(w in cmd for w in ["my homework", "do my homework",
                                              "help with homework", "solve my homework"])
    wants_solve    = (not is_homework_cmd) and any(w in cmd for w in [
                        "solve", "answer", "calculate", "find",
                        "scan", "read", "check", "look at"])
    wants_explain  = any(w in cmd for w in ["explain", "teach", "understand",
                                             "concept", "why", "how does", "tell me about"])
    wants_quiz     = any(w in cmd for w in ["quiz", "test me", "practice", "questions"])
    wants_simulate = any(w in cmd for w in ["simulate", "animation", "show me",
                                             "demonstrate", "animate"])
    wants_graph    = any(w in cmd for w in ["graph", "plot", "chart", "visualize"])
    wants_similar  = any(w in cmd for w in ["similar", "more problems", "practice problems",
                                             "more questions", "like this"])
    wants_step     = any(w in cmd for w in ["step by step", "walk me through",
                                             "show working", "full working"])
    wants_wrong    = any(w in cmd for w in ["explain wrong", "wrong answers",
                                             "what i got wrong", "my mistakes"])

    print(f"[CHAIN] screen:{wants_screen} solve:{wants_solve} explain:{wants_explain} "
          f"quiz:{wants_quiz} sim:{wants_simulate} graph:{wants_graph}")

    if wants_screen and (wants_solve or wants_explain or wants_quiz
                         or wants_simulate or wants_graph):
        chains.append("screen_capture")

    # always solve when screen is involved — even if user only said "explain"
    # because we need to read the screen content first
    if wants_solve or wants_screen:
        chains.append("solve")

    if wants_explain:
        chains.append("explain")

    if wants_quiz:
        chains.append("quiz")

    if wants_simulate:
        chains.append("simulate")

    if wants_graph:
        chains.append("graph")

    if wants_similar:
        chains.append("similar_problems")

    if wants_step:
        chains.append("step_by_step")

    if wants_wrong:
        chains.append("explain_wrong")

    seen   = set()
    unique = []
    for c in chains:
        if c not in seen:
            seen.add(c)
            unique.append(c)

    print(f"[CHAIN] Final: {unique if unique else 'None'}")
    return unique if unique else None

# ================= 🎤 VOICE INPUT =================

def listen_for_prompt(timeout=8, phrase_time_limit=15):
    """
    Listen for a voice response from the user.
    Returns transcribed text or empty string on failure.
    """
    r = sr.Recognizer()
    try:
        with sr.Microphone() as source:
            r.adjust_for_ambient_noise(source, duration=0.5)
            print("[LISTEN] Waiting for user response...")
            audio = r.listen(
                source,
                timeout=timeout,
                phrase_time_limit=phrase_time_limit
            )
        text = r.recognize_google(audio).lower().strip()
        print(f"[LISTEN] Heard: '{text}'")
        return text
    except sr.WaitTimeoutError:
        print("[LISTEN] Timeout — no speech detected.")
        return ""
    except sr.UnknownValueError:
        print("[LISTEN] Could not understand audio.")
        return ""
    except Exception as e:
        print(f"[LISTEN] Error: {e}")
        return ""

# ================= 🤖 AI CALLS =================

def _thinking_thread(stop_event):
    import threading
    messages = [
        "Still working on it. This might take a moment.",
        "Almost there. The AI is still processing.",
        "Still thinking. Complex problems take a little longer.",
        "Bear with me. Nearly done.",
    ]
    idx = 0
    while not stop_event.wait(30):
        speak(messages[idx % len(messages)])
        idx += 1


def ask_ai(prompt, max_tokens=2048):
    """
    Text AI tasks — routes through ai_router.
    Simple/local queries → Ollama; complex/research queries → Groq.
    Vision tasks (ask_ai_with_image) always stay on Groq.
    """
    import threading
    print(f"[STUDY AI] Routing text request via ai_router")
    print(f"[STUDY AI] Prompt preview: {prompt[:100].strip()}...")

    stop_event = threading.Event()
    t = threading.Thread(target=_thinking_thread, args=(stop_event,), daemon=True)
    t.start()

    try:
        result = _routed_ask_ai(prompt, command=prompt, has_live_data=False)
        print(f"[STUDY AI] Response: {len(result)} chars")
        return result
    except Exception as e:
        print(f"[STUDY AI] Error: {e}")
        return ""
    finally:
        stop_event.set()


def ask_ai_with_image(prompt, image_path):
    """Image + vision tasks go through Groq vision model."""
    import threading
    print(f"[GROQ VISION] Sending image | model: {GROQ_VISION_MODEL}")
    print(f"[GROQ VISION] Image: {image_path}")

    if not groq_client or not os.path.exists(image_path):
        print("[GROQ VISION] Client not ready or image missing.")
        return ""

    stop_event = threading.Event()
    t = threading.Thread(target=_thinking_thread, args=(stop_event,), daemon=True)
    t.start()

    start = datetime.now()
    try:
        with open(image_path, "rb") as f:
            image_b64 = base64.b64encode(f.read()).decode("utf-8")

        ext      = Path(image_path).suffix.lower()
        mime_map = {".png": "image/png", ".jpg": "image/jpeg",
                    ".jpeg": "image/jpeg", ".webp": "image/webp"}
        mime     = mime_map.get(ext, "image/png")

        response = groq_client.chat.completions.create(
            model=GROQ_VISION_MODEL,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image_url",
                     "image_url": {"url": f"data:{mime};base64,{image_b64}"}},
                    {"type": "text", "text": prompt}
                ]
            }],
            max_tokens=2048,
        )
        elapsed = (datetime.now() - start).seconds
        result  = response.choices[0].message.content.strip()
        print(f"[GROQ VISION] Response in {elapsed}s | {len(result)} chars")
        return result
    except Exception as e:
        print(f"[GROQ VISION] Error: {e}")
        return ""
    finally:
        stop_event.set()


def ask_gemini(prompt, max_tokens=3000):
    """Code generation tasks (graphs, simulations) — uses Groq llama-3.3-70b."""
    import threading
    print(f"[GROQ CODE] Sending code generation request")
    print(f"[GROQ CODE] Prompt preview: {prompt[:100].strip()}...")

    if not groq_client:
        print("[GROQ CODE] Client not initialized.")
        return ""

    stop_event = threading.Event()
    t = threading.Thread(target=_thinking_thread, args=(stop_event,), daemon=True)
    t.start()

    start = datetime.now()
    try:
        response = groq_client.chat.completions.create(
            model=GROQ_TEXT_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an expert JavaScript and HTML developer. "
                        "When asked to generate code, return ONLY the raw HTML/JS snippet. "
                        "No markdown fences, no explanations outside the code."
                    )
                },
                {"role": "user", "content": prompt}
            ],
            max_tokens=3000,
        )
        elapsed = (datetime.now() - start).seconds
        result  = response.choices[0].message.content.strip()
        # strip markdown fences if model added them
        result  = re.sub(r"```html|```javascript|```js|```", "", result).strip()
        print(f"[GROQ CODE] Response in {elapsed}s | {len(result)} chars")
        return result
    except Exception as e:
        print(f"[GROQ CODE] Error: {e}")
        return ""
    finally:
        stop_event.set()

# ================= 📸 SCREEN CAPTURE =================

def take_screenshot():
    path = os.path.join(OUTPUT_DIR, "screen_snapshot.png")
    try:
        screenshot = pyautogui.screenshot()
        screenshot.save(path)
        size = os.path.getsize(path)
        print(f"[SCREENSHOT] Saved: {path} ({size / 1024:.1f} KB)")
        SESSION["last_image"] = path
    except Exception as e:
        print(f"[SCREENSHOT] Failed: {e}")
        path = None
    return path


def take_screenshot_with_prompt():
    """
    Ask the user what they want before taking the screenshot.
    Returns (image_path, user_prompt).
    Only used in study mode.
    """
    speak("What would you like me to help you with from your screen? For example: solve this question, or explain this diagram.")
    user_prompt = listen_for_prompt()

    if not user_prompt:
        speak("I did not catch that. I will go ahead and solve whatever is on screen.")
        user_prompt = "Solve the question shown on screen and explain it step by step."

    print(f"[SCREEN PROMPT] User asked: '{user_prompt}'")
    speak("Got it. Scanning your screen now.")

    image_path = take_screenshot()
    return image_path, user_prompt

# ================= 🌐 BROWSER OUTPUT =================

def open_in_browser(html_content, title="FRIDAY Study"):
    filename = os.path.join(
        OUTPUT_DIR,
        title.replace(" ", "_") + "_" + datetime.now().strftime("%H%M%S") + ".html"
    )
    try:
        with open(filename, "w", encoding="utf-8") as f:
            f.write(html_content)
        print(f"[BROWSER] Saved: {filename} ({len(html_content) / 1024:.1f} KB)")
        webbrowser.open("file:///" + filename.replace("\\", "/"))
    except Exception as e:
        print(f"[BROWSER] Failed: {e}")
    return filename


# ================= 🧹 LATEX CLEANUP =================

def clean_latex(text):
    """
    Fix common malformed LaTeX patterns from AI output
    before injecting into HTML so MathJax renders correctly.
    """
    import re as _re

    # Fix $/random/$ -> $random$ (remove stray slashes around dollar signs)
    text = _re.sub(r'\$/([^$]+)/\$', r'$\1$', text)

    # Fix $$ used as inline (single line) -> keep as display
    # Fix lone $ signs that arent math (e.g. dollar amounts like $5)
    # Only treat as math if followed by letters/symbols not spaces+digits
    # (heuristic: $5 or $10 are money, $x or $\frac are math)
    def fix_dollar(m):
        inner = m.group(1)
        # if it looks like a number only -> not math, escape it
        if _re.match(r'^\d+(\.\d+)?$', inner.strip()):
            return m.group(0).replace('$', '&#36;')
        return m.group(0)

    # Fix \( and \) that got double-escaped as \\( \\)
    text = text.replace('\\\\(', '\\(').replace('\\\\)', '\\)')
    text = text.replace('\\\\[', '\\[').replace('\\\\]', '\\]')

    # Fix ** bold markdown -> <strong>
    text = _re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)

    # Fix * italic markdown -> <em>
    text = _re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)', r'<em>\1</em>', text)

    # Fix ### headings from AI
    text = _re.sub(r'^###\s+(.+)$', r'<h3>\1</h3>', text, flags=_re.MULTILINE)
    text = _re.sub(r'^##\s+(.+)$',  r'<h3>\1</h3>', text, flags=_re.MULTILINE)

    # Fix numbered list items 1. 2. 3. -> styled divs
    text = _re.sub(
        r'^(\d+)\.\s+(.+)$',
        r'<div style="margin:6px 0; padding-left:20px;"><span style="color:#7eb8ff; font-weight:bold;">\1.</span> \2</div>',
        text, flags=_re.MULTILINE
    )

    return text

def wrap_html(title, body, extra_scripts=""):
    return f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="UTF-8">
  <title>{title}</title>
  <!-- MathJax with proper config for $...$ and $$...$$ rendering -->
  <script>
  window.MathJax = {{
    tex: {{
      inlineMath: [['$', '$'], ['\\(', '\\)']],
      displayMath: [['$$', '$$'], ['\\[', '\\]']],
      processEscapes: true,
      tags: 'ams'
    }},
    options: {{ skipHtmlTags: ['script','noscript','style','textarea','pre'] }},
    startup: {{
      ready() {{
        MathJax.startup.defaultReady();
        MathJax.startup.promise.then(() => {{
          MathJax.typesetPromise();
        }});
      }}
    }}
  }};
  </script>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/mathjax/3.2.2/es5/tex-mml-chtml.min.js"></script>
  <!-- Plotly for Desmos-style interactive graphs -->
  <script src="https://cdn.plot.ly/plotly-2.27.0.min.js"></script>
  <!-- p5.js for smooth canvas animations -->
  <script src="https://cdnjs.cloudflare.com/ajax/libs/p5.js/1.9.0/p5.min.js"></script>
  <!-- Matter.js for physics simulations -->
  <script src="https://cdnjs.cloudflare.com/ajax/libs/matter-js/0.19.0/matter.min.js"></script>
  <style>
    body {{
      font-family: 'Segoe UI', sans-serif;
      background: #0f0f1a;
      color: #e0e0ff;
      max-width: 960px;
      margin: 40px auto;
      padding: 30px;
      line-height: 1.8;
    }}
    h1 {{ color: #7eb8ff; border-bottom: 1px solid #2a2a4e; padding-bottom: 10px; }}
    h2 {{ color: #a0d0ff; margin-top: 30px; }}
    h3 {{ color: #c0e0ff; }}
    .step {{
      background: #1a1a2e;
      border-left: 4px solid #7eb8ff;
      padding: 15px 20px;
      margin: 15px 0;
      border-radius: 4px;
    }}
    .answer {{
      background: #1a2e1a;
      border-left: 4px solid #7eff7e;
      padding: 15px 20px;
      margin: 15px 0;
      border-radius: 4px;
    }}
    .warning {{
      background: #2e1a1a;
      border-left: 4px solid #ff7e7e;
      padding: 15px 20px;
      margin: 15px 0;
      border-radius: 4px;
    }}
    .wrong-review-card {{
      background: #1a1020;
      border: 1px solid #4a2a4a;
      border-radius: 10px;
      padding: 20px 25px;
      margin: 20px 0;
    }}
    .wrong-topic-label {{
      color: #ff9eb8;
      font-size: 11px;
      font-weight: bold;
      letter-spacing: 2px;
      text-transform: uppercase;
      margin-bottom: 12px;
      padding-bottom: 8px;
      border-bottom: 1px solid #4a2a4a;
    }}
    .wrong-explanation {{
      color: #e0e0ff;
      line-height: 1.9;
    }}
    .sim-container {{
      width: 100%;
      position: relative;
      background: #080810;
      border-radius: 10px;
      overflow: hidden;
      padding: 0;
      margin: 15px 0;
    }}
    .sim-controls {{
      background: #13131f;
      border-top: 1px solid #2a2a4e;
      padding: 15px 20px;
      display: flex;
      flex-wrap: wrap;
      gap: 15px;
      align-items: center;
    }}
    .sim-controls label {{
      color: #a0d0ff;
      font-size: 13px;
      display: flex;
      align-items: center;
      gap: 8px;
    }}
    .sim-controls input[type=range] {{
      width: 120px;
      accent-color: #7eb8ff;
    }}
    .sim-controls button {{
      background: #2a2a4e;
      color: #e0e0ff;
      border: 1px solid #7eb8ff;
      padding: 6px 16px;
      border-radius: 4px;
      cursor: pointer;
      font-size: 13px;
    }}
    .sim-controls button:hover {{ background: #3a3a6e; }}
    .sim-stats {{
      background: #0a0a18;
      border-top: 1px solid #1a1a3e;
      padding: 10px 20px;
      display: flex;
      flex-wrap: wrap;
      gap: 20px;
      font-size: 13px;
      color: #7eff7e;
      font-family: monospace;
    }}
    .section {{
      background: #13131f;
      border: 1px solid #2a2a4e;
      border-radius: 10px;
      padding: 20px 25px;
      margin: 25px 0;
    }}
    .section-title {{
      color: #7eb8ff;
      font-size: 13px;
      font-weight: bold;
      letter-spacing: 2px;
      text-transform: uppercase;
      margin-bottom: 15px;
    }}
    .quiz-question {{
      background: #1a1a2e;
      border: 1px solid #2a2a4e;
      padding: 20px;
      margin: 15px 0;
      border-radius: 8px;
    }}
    .quiz-question button {{
      background: #2a2a4e;
      color: #e0e0ff;
      border: 1px solid #4a4a7e;
      padding: 8px 16px;
      margin: 5px;
      border-radius: 4px;
      cursor: pointer;
      font-size: 14px;
      transition: background 0.2s;
    }}
    .quiz-question button:hover {{ background: #3a3a6e; }}
    .correct {{ background: #1a3a1a !important; border-color: #7eff7e !important; color: #7eff7e !important; }}
    .wrong   {{ background: #3a1a1a !important; border-color: #ff7e7e !important; color: #ff7e7e !important; }}
    .explanation-box {{
      display: none;
      margin-top: 12px;
      padding: 12px 16px;
      background: #0f1f2f;
      border-left: 3px solid #7eb8ff;
      border-radius: 4px;
      font-size: 14px;
      color: #a0d0ff;
    }}
    canvas {{ border: 1px solid #2a2a4e; border-radius: 8px; margin: 20px 0; display: block; }}
    pre {{
      background: #111;
      padding: 15px;
      border-radius: 6px;
      overflow-x: auto;
      color: #a0ffa0;
      font-size: 13px;
    }}
    .tag {{
      display: inline-block;
      background: #2a2a4e;
      color: #7eb8ff;
      padding: 3px 12px;
      border-radius: 12px;
      font-size: 12px;
      margin-bottom: 15px;
    }}
    .score-board {{
      font-size: 20px;
      color: #7eb8ff;
      margin: 15px 0;
      padding: 10px 20px;
      background: #1a1a2e;
      border-radius: 8px;
      display: inline-block;
    }}
    hr {{ border-color: #2a2a4e; margin: 30px 0; }}
  </style>
  {extra_scripts}
</head>
<body>
  <span class="tag">FRIDAY Study — {BOARD} Grade {GRADE}</span>
  <h1>{title}</h1>
  {body}
</body>
</html>"""

# ================= 🧩 ATOMIC FEATURE FUNCTIONS =================

def _build_solution(topic_or_image, from_screen=False, extra_instruction="",
                    user_screen_prompt=None):
    """Core solver. Returns (html, detected_topic)."""

    if from_screen or topic_or_image == "__screen__":
        image_path = SESSION.get("last_image") or take_screenshot()
        print(f"[SOLVE] Screen mode | image: {image_path}")

        # use user's specific prompt if provided, else generic solve prompt
        if user_screen_prompt:
            vision_prompt = f"""You are a {BOARD} Grade {GRADE} tutor for Maths, Physics, and Chemistry.

The student asks: "{user_screen_prompt}"

Looking at their screen:
1. Address exactly what the student asked.
2. If it is a question/problem, solve it step by step with full working.
3. State the final answer clearly.
4. Explain the concept behind it.
5. Mention the CBSE chapter/topic.
{extra_instruction}

Use LaTeX: inline \\( \\), display \\[ \\]."""
        else:
            vision_prompt = f"""You are a {BOARD} Grade {GRADE} tutor for Maths, Physics, and Chemistry.

A student has shared their screen.
1. State the question/problem exactly as you see it.
2. Identify the CBSE chapter/topic.
3. Solve it fully with step-by-step working.
4. State the final answer clearly.
5. Briefly explain the concept.
{extra_instruction}

Use LaTeX: inline \\( \\), display \\[ \\]."""

        result       = ask_ai_with_image(vision_prompt, image_path)
        source_label = "Screen Capture"

    else:
        print(f"[SOLVE] Text mode | problem: '{topic_or_image[:80]}'")
        prompt = f"""You are a {BOARD} Grade {GRADE} tutor.

Solve: "{topic_or_image}"

1. Restate the problem clearly.
2. Identify subject and CBSE chapter/topic.
3. Solve step-by-step with full working.
4. State final answer.
5. Briefly explain the concept.
{extra_instruction}

Use LaTeX: inline \\( \\), display \\[ \\]."""
        result       = ask_ai(prompt)
        source_label = "Problem"

    if not result:
        print("[SOLVE] Empty result from AI.")
        result = "Could not generate a solution. Please try again."

    # detect topic
    topic_prompt   = f'In one short phrase, what is the main topic of this solution?\n\n"{result[:300]}"\n\nReturn only the topic name, nothing else.'
    detected_topic = ask_ai(topic_prompt) or topic_or_image
    print(f"[SOLVE] Detected topic: '{detected_topic}'")

    SESSION["last_solution"] = result
    SESSION["last_topic"]    = detected_topic

    html = f"""
    <div class="section">
      <div class="section-title">📋 {source_label}</div>
      <div class="step">{clean_latex(result).replace(chr(10), "<br>")}</div>
    </div>"""

    return html, detected_topic


def _build_explanation(topic, subject=None):
    if not subject:
        subject = detect_subject(topic)

    print(f"[EXPLAIN] topic: '{topic}' | subject: '{subject}'")

    SESSION["last_topic"]   = topic
    SESSION["last_subject"] = subject

    prompt = f"""You are a {BOARD} Grade {GRADE} {subject} teacher.

Explain: "{topic}"

Structure your response with these sections:
1. Simple definition
2. Core concept in depth
3. Step-by-step derivation (if applicable)
4. Real-world example or analogy
5. Common student mistakes
6. CBSE exam tips

Use LaTeX: inline \\( \\), display \\[ \\].
Be thorough and student-friendly."""

    result = ask_ai(prompt)

    if not result:
        print("[EXPLAIN] Empty result.")
        result = "Could not generate explanation. Please try again."

    html = f"""
    <div class="section">
      <div class="section-title">📖 Concept Explanation</div>
      <div class="step">{clean_latex(result).replace(chr(10), "<br>")}</div>
    </div>"""

    return html


def parse_quiz_params(command):
    """Extract duration (minutes) and question count from command."""
    cmd = command.lower()

    # extract question count: "20 questions", "10 mcqs"
    q_match = re.search(r"(\d+)\s*(?:questions?|mcqs?|problems?|qs?)", cmd)
    num_questions = int(q_match.group(1)) if q_match else 5
    num_questions = min(max(num_questions, 3), 30)  # clamp 3-30

    # extract duration: "10 minutes", "half an hour", "30 mins"
    t_match = re.search(r"(\d+)\s*(?:minutes?|mins?)", cmd)
    if t_match:
        duration_minutes = int(t_match.group(1))
    elif "half an hour" in cmd or "30 min" in cmd:
        duration_minutes = 30
    elif "one hour" in cmd or "1 hour" in cmd or "an hour" in cmd:
        duration_minutes = 60
    else:
        duration_minutes = None  # no timer

    print(f"[QUIZ PARAMS] questions={num_questions} | duration={duration_minutes}m")
    return num_questions, duration_minutes


def _build_quiz(topic, subject=None, num_questions=5, duration_minutes=None):
    if not subject:
        subject = detect_subject(topic)

    print(f"[QUIZ] topic: '{topic}' | subject: '{subject}' | n: {num_questions} | timer: {duration_minutes}m")

    prompt = f"""Generate {num_questions} high-quality MCQ questions for Grade {GRADE} CBSE {subject} on: "{topic}"

Use real CBSE exam-style questions. Include numerical problems, conceptual questions, and application-based questions.

Return ONLY a JSON array. Each object must have:
- "question": question text (no LaTeX, use plain text math)
- "options": array of 4 strings starting with A), B), C), D)
- "answer": correct letter (A/B/C/D)
- "explanation": detailed explanation of why this answer is correct
- "topic": specific sub-topic this tests (e.g. "Laws of Motion - Newton\'s Second Law")

Return ONLY valid JSON array, no markdown, no extra text."""

    result = ask_ai(prompt)

    if not result:
        print("[QUIZ] Empty result.")
        return "<p>Could not generate quiz. Please try again.</p>", []

    try:
        clean     = re.sub(r"```json|```", "", result).strip()
        questions = json.loads(clean)
        print(f"[QUIZ] Parsed {len(questions)} questions.")
    except Exception as e:
        print(f"[QUIZ] JSON parse failed: {e} | raw: {result[:200]}")
        return "<p>Could not generate quiz. Please try again.</p>", []

    SESSION["last_topic"]   = topic
    SESSION["last_subject"] = subject

    results_data = json.dumps([
        {"topic": q.get("topic", topic), "explanation": q.get("explanation", "")}
        for q in questions
    ])

    # path where browser saves wrong topics so Python can read them
    wrong_json_path = os.path.join(OUTPUT_DIR, "wrong_topics.json").replace("\\", "/").replace("\\\\", "/")

    # timer HTML
    timer_html = ""
    timer_script = ""
    if duration_minutes:
        total_seconds = duration_minutes * 60
        timer_html = f"""
        <div id="timerBox" style="position:fixed; top:20px; right:20px; background:#1a1a2e;
             border:2px solid #7eb8ff; border-radius:10px; padding:15px 25px;
             font-size:24px; color:#7eb8ff; font-weight:bold; z-index:999;">
          ⏱ <span id="timerDisplay">{duration_minutes}:00</span>
        </div>"""
        timer_script = f"""
        // ── Timer ──
        let timeLeft = {total_seconds};
        const timerDisplay = document.getElementById('timerDisplay');
        const timerInterval = setInterval(() => {{
            timeLeft--;
            const m = Math.floor(timeLeft / 60);
            const s = timeLeft % 60;
            timerDisplay.textContent = m + ':' + String(s).padStart(2, '0');
            if (timeLeft <= 60) timerDisplay.style.color = '#ff7e7e';
            if (timeLeft <= 0) {{
                clearInterval(timerInterval);
                timerDisplay.textContent = "TIME UP";
                showFinalResults({{ score: score, total: {len(questions)}, wrongTopics: wrongTopics }});
            }}
        }}, 1000);"""

    quiz_html = f"""
    {timer_html}
    <div class="score-board" id="scoreBoard">Score: 0 / {len(questions)}</div>"""

    for i, q in enumerate(questions):
        opts    = q.get("options", [])
        correct = q.get("answer", "A")
        qtopic  = q.get("topic", topic).replace("'", " ").replace('"', " ")

        options_html = ""
        for opt in opts:
            letter = opt[0] if opt else "A"
            options_html += f'<button onclick="checkAnswer(this,\'{letter}\',\'{correct}\',\'exp{i}\',\'{qtopic}\')">{{opt}}</button>\n'

        quiz_html += f"""
        <div class="quiz-question" id="qq{i}">
            <strong>Q{i+1}. {q.get("question","")}</strong><br><br>
            {options_html}
            <div class="explanation-box" id="exp{i}"></div>
        </div>"""

    quiz_script = f"""
    <script>
    let score = 0;
    let answered = {{}};
    let wrongTopics = [];
    const questionsData = {results_data};
    const WRONG_JSON_PATH = "{wrong_json_path}";
    {timer_script}

    function checkAnswer(btn, selected, correct, expId, qtopic) {{
        const qDiv = btn.parentElement;
        if (answered[qDiv.id]) return;
        answered[qDiv.id] = true;
        const expDiv = document.getElementById(expId);
        const idx = parseInt(expId.replace('exp',''));
        if (selected === correct) {{
            btn.classList.add('correct');
            score++;
            document.getElementById('scoreBoard').textContent = 'Score: ' + score + ' / {len(questions)}';
            expDiv.style.display = 'block';
            expDiv.innerHTML = '<strong>Correct!</strong> ' + questionsData[idx].explanation;
            expDiv.style.borderColor = '#7eff7e';
        }} else {{
            btn.classList.add('wrong');
            wrongTopics.push(qtopic);
            const buttons = qDiv.querySelectorAll('button');
            buttons.forEach(b => {{ if (b.textContent.startsWith(correct)) b.classList.add('correct'); }});
            expDiv.style.display = 'block';
            expDiv.innerHTML = '<strong>Incorrect.</strong> ' + questionsData[idx].explanation;
            expDiv.style.borderColor = '#ff7e7e';
        }}
        if (Object.keys(answered).length === {len(questions)}) {{
            showFinalResults({{ score: score, total: {len(questions)}, wrongTopics: wrongTopics }});
        }}
    }}

    function saveWrongTopics(wrongTopics) {{
        // Save wrong topics to a file FRIDAY can read
        const data = JSON.stringify({{
            wrongTopics: [...new Set(wrongTopics)],
            timestamp: new Date().toISOString(),
            topic: "{topic}",
            score: score,
            total: {len(questions)}
        }});
        // Write via fetch to a local file server isn't possible directly,
        // so we use a hidden link download trick
        const blob = new Blob([data], {{type: "application/json"}});
        const url  = URL.createObjectURL(blob);
        const a    = document.createElement("a");
        a.href     = url;
        a.download = "wrong_topics.json";
        a.style.display = "none";
        document.body.appendChild(a);
        a.click();
        URL.revokeObjectURL(url);
        console.log("Wrong topics saved:", wrongTopics);
    }}

    function showFinalResults(results) {{
        if (typeof timerInterval !== "undefined") clearInterval(timerInterval);
        const div = document.createElement("div");
        div.className = "section";
        const uniqueWrong = [...new Set(results.wrongTopics)];
        div.innerHTML = "<div class=\"section-title\">🎯 Final Results</div>" +
            "<p>You scored <strong>" + results.score + " / {len(questions)}</strong></p>" +
            (uniqueWrong.length > 0
                ? "<p>Topics to review: <strong>" + uniqueWrong.join(", ") + "</strong></p>" +
                  "<p style=\"color:#a0d0ff\">Say <em>explain my wrong answers</em> to FRIDAY.</p>" +
                  "<button onclick=\"saveAndReview()\" style=\"background:#2a2a4e;color:#7eb8ff;border:1px solid #7eb8ff;padding:10px 20px;border-radius:6px;cursor:pointer;font-size:15px;margin-top:10px;\">💾 Save Wrong Answers for FRIDAY</button>"
                : "<p style=\"color:#7eff7e\">Perfect score! Excellent work.</p>");
        document.body.appendChild(div);
        // auto-save wrong topics
        if (uniqueWrong.length > 0) saveWrongTopics(uniqueWrong);
    }}

    function saveAndReview() {{
        saveWrongTopics([...new Set(wrongTopics)]);
        alert("Wrong answers saved! Now say \"explain my wrong answers\" to FRIDAY.");
    }}
    </script>"""

    html = f"""
    <div class="section">
      <div class="section-title">🧪 Quiz: {topic}{f" ({duration_minutes} min)" if duration_minutes else ""}</div>
      {quiz_html}
      {quiz_script}
    </div>"""

    return html, questions
    if not subject:
        subject = detect_subject(topic)

    print(f"[QUIZ] topic: '{topic}' | subject: '{subject}' | n: {num_questions}")

    prompt = f"""Generate {num_questions} MCQ questions for Grade {GRADE} CBSE {subject} on: "{topic}"

Return ONLY a JSON array. Each object must have:
- "question": question text (no LaTeX)
- "options": array of 4 strings starting with A), B), C), D)
- "answer": correct letter (A/B/C/D)
- "explanation": why this answer is correct
- "topic": specific sub-topic this tests

Return ONLY valid JSON array, no markdown, no extra text."""

    result = ask_ai(prompt)

    if not result:
        print("[QUIZ] Empty result.")
        return "<p>Could not generate quiz. Please try again.</p>", []

    try:
        clean     = re.sub(r"```json|```", "", result).strip()
        questions = json.loads(clean)
        print(f"[QUIZ] Parsed {len(questions)} questions.")
    except Exception as e:
        print(f"[QUIZ] JSON parse failed: {e} | raw: {result[:200]}")
        return "<p>Could not generate quiz. Please try again.</p>", []

    SESSION["last_topic"]   = topic
    SESSION["last_subject"] = subject

    results_data = json.dumps([
        {"topic": q.get("topic", topic), "explanation": q.get("explanation", "")}
        for q in questions
    ])

    quiz_html = f'<div class="score-board" id="scoreBoard">Score: 0 / {len(questions)}</div>'

    for i, q in enumerate(questions):
        opts    = q.get("options", [])
        correct = q.get("answer", "A")
        qtopic  = q.get("topic", topic).replace("'", " ").replace('"', " ")

        options_html = ""
        for opt in opts:
            letter = opt[0] if opt else "A"
            options_html += f'<button onclick="checkAnswer(this,\'{letter}\',\'{correct}\',\'exp{i}\',\'{qtopic}\')">{opt}</button>\n'

        quiz_html += f"""
        <div class="quiz-question" id="qq{i}">
            <strong>Q{i+1}. {q.get("question","")}</strong><br><br>
            {options_html}
            <div class="explanation-box" id="exp{i}"></div>
        </div>"""

    quiz_script = f"""
    <script>
    let score = 0;
    let answered = {{}};
    let wrongTopics = [];
    const questionsData = {results_data};

    function checkAnswer(btn, selected, correct, expId, qtopic) {{
        const qDiv = btn.parentElement;
        if (answered[qDiv.id]) return;
        answered[qDiv.id] = true;
        const expDiv = document.getElementById(expId);
        const idx = parseInt(expId.replace('exp',''));
        if (selected === correct) {{
            btn.classList.add('correct');
            score++;
            document.getElementById('scoreBoard').textContent = 'Score: ' + score + ' / {len(questions)}';
            expDiv.style.display = 'block';
            expDiv.innerHTML = '<strong>Correct!</strong> ' + questionsData[idx].explanation;
            expDiv.style.borderColor = '#7eff7e';
        }} else {{
            btn.classList.add('wrong');
            wrongTopics.push(qtopic);
            const buttons = qDiv.querySelectorAll('button');
            buttons.forEach(b => {{ if (b.textContent.startsWith(correct)) b.classList.add('correct'); }});
            expDiv.style.display = 'block';
            expDiv.innerHTML = '<strong>Incorrect.</strong> ' + questionsData[idx].explanation;
            expDiv.style.borderColor = '#ff7e7e';
        }}
        if (Object.keys(answered).length === {len(questions)}) {{
            showFinalResults({{ score: score, total: {len(questions)}, wrongTopics: wrongTopics }});
        }}
    }}
    function showFinalResults(results) {{
        const div = document.createElement('div');
        div.className = 'section';
        div.innerHTML = '<div class="section-title">Final Results</div>' +
            '<p>You scored <strong>' + results.score + ' / {len(questions)}</strong></p>' +
            (results.wrongTopics.length > 0
                ? '<p>Topics to review: <strong>' + [...new Set(results.wrongTopics)].join(', ') + '</strong></p>' +
                  '<p style="color:#a0d0ff">Say explain my wrong answers to FRIDAY for detailed explanations.</p>'
                : '<p style="color:#7eff7e">Perfect score! Great work.</p>');
        document.body.appendChild(div);
    }}
    </script>"""

    html = f"""
    <div class="section">
      <div class="section-title">🧪 Quiz: {topic}</div>
      {quiz_html}
      {quiz_script}
    </div>"""

    return html, questions


def _build_simulation(topic):
    """Uses Gemini for reliable code generation."""
    print(f"[SIMULATE] Generating simulation for: '{topic}' via Gemini")

    prompt = f"""Write a COMPLETE self-contained HTML snippet (div and script tags only, NO html/head/body tags) for an interactive science simulation.

Topic: "{topic}"
Student: Grade {GRADE} CBSE

Requirements:
- Canvas-based animation using vanilla JavaScript only
- Interactive sliders or buttons to control parameters
- Real-time display of values (e.g. velocity, time, angle)
- Dark theme: background #0f0f1a, accent colors #7eb8ff and #7eff7e
- Physically accurate for CBSE Grade {GRADE} syllabus
- A title label inside the simulation showing what is being simulated

Examples by topic:
- Projectile: show trajectory arc, update range and max height live
- SHM/Pendulum: animate bob, show period and energy
- Wave: show transverse wave, control amplitude and frequency
- Circuit: animate current flow, show voltage/resistance/current values
- Optics: show ray diagram with movable lens

Make it visually impressive. Return ONLY the HTML snippet, no markdown."""

    result = ask_gemini(prompt)

    if not result or len(result) < 100:
        print("[SIMULATE] Empty or too short result.")
        return "<p>Could not generate simulation. Try being more specific about the topic.</p>"

    # strip markdown fences if Gemini added them
    result = re.sub(r"```html|```", "", result).strip()

    print(f"[SIMULATE] Code received: {len(result)} chars")

    html = f"""
    <div class="section">
      <div class="section-title">🔬 Simulation: {topic}</div>
      {result}
    </div>"""

    return html


def _build_graph(topic):
    """Plotly.js interactive graphs — Desmos-style with zoom/pan and working sliders."""
    print(f"[GRAPH] Generating Plotly graph for: '{topic}'")

    prompt = f"""A Grade {GRADE} CBSE student wants an interactive Plotly.js graph for: "{topic}"

Write a self-contained HTML snippet (div and script only, NO html/head/body tags).
Plotly.js 2.27.0 is already loaded on the page.

STRICT REQUIREMENTS:
1. Use Plotly.newPlot() — NOT Chart.js
2. Compute data mathematically in JavaScript using Math functions
3. Generate at least 300 points for smooth curves
4. Dark theme: {{plot_bgcolor:'#0d0d1a', paper_bgcolor:'transparent', font:{{color:'#e0e0ff'}}, xaxis:{{gridcolor:'#1e1e3a', zerolinecolor:'#3a3a6e'}}, yaxis:{{gridcolor:'#1e1e3a', zerolinecolor:'#3a3a6e'}}}}
5. Enable zoom, pan, hover (Plotly default — just set dragmode:'zoom' in layout)
6. Unique div id: "plotly_{abs(hash(topic)) % 99999}"
7. Div style: width:100%; height:480px;
8. If the topic suits multiple curves (e.g. different initial conditions), plot all of them with a legend

WORKING SLIDERS — this is mandatory:
- Add HTML range sliders BELOW the graph div for key parameters (e.g. initial velocity, angle, amplitude)
- Each slider must have an oninput handler that RECOMPUTES the data and calls Plotly.react() to update the graph
- Show the current slider value next to each slider label

Example slider pattern:
<div style="padding:12px; background:#13131f; border-top:1px solid #2a2a4e; display:flex; flex-wrap:wrap; gap:15px; align-items:center;">
  <label style="color:#a0d0ff; font-size:13px;">Param: <span id="paramVal">5</span>
    <input type="range" min="1" max="20" value="5" style="width:130px; accent-color:#7eb8ff; margin-left:8px;"
      oninput="document.getElementById('paramVal').textContent=this.value; updateGraph(+this.value);">
  </label>
</div>
<script>
function updateGraph(param) {{
  // recompute x and y arrays
  const x = [], y = [];
  for (let i = 0; i <= 300; i++) {{
    const t = i / 30;
    x.push(t);
    y.push(param * t); // example
  }}
  Plotly.react('plotly_{abs(hash(topic)) % 99999}', [{{x, y, ...traceStyle}}], layout);
}}
</script>

Include a brief paragraph BEFORE the graph explaining what is shown.
Return ONLY the HTML snippet, no markdown."""

    result = ask_gemini(prompt)

    if not result or len(result) < 100:
        print("[GRAPH] Empty result.")
        return "<p>Could not generate graph. Please try again.</p>"

    result = re.sub(r"```html|```javascript|```js|```", "", result).strip()
    print(f"[GRAPH] Code received: {len(result)} chars")

    html = f"""
    <div class="section">
      <div class="section-title">📊 Graph: {topic}</div>
      {result}
    </div>"""

    return html


def _build_wrong_answer_explanations():
    """
    Read wrong topics from the downloaded JSON file.
    The browser saves wrong_topics.json to the user's Downloads folder.
    We search Downloads + OUTPUT_DIR for it.
    """
    import glob

    wrong_topics = []
    json_data    = None

    # search locations: Downloads folder + FRIDAY temp folder
    search_paths = [
        os.path.join(os.path.expanduser("~"), "Downloads", "wrong_topics.json"),
        os.path.join(os.path.expanduser("~"), "Desktop",   "wrong_topics.json"),
        os.path.join(OUTPUT_DIR, "wrong_topics.json"),
    ]

    for path in search_paths:
        if os.path.exists(path):
            try:
                with open(path, "r") as f:
                    json_data = json.load(f)
                wrong_topics = json_data.get("wrongTopics", [])
                print(f"[WRONG] Loaded from: {path} | topics: {wrong_topics}")
                break
            except Exception as e:
                print(f"[WRONG] Failed to read {path}: {e}")

    # also check session as fallback
    if not wrong_topics:
        wrong_topics = SESSION.get("last_wrong_topics", [])
        print(f"[WRONG] Falling back to session topics: {wrong_topics}")

    if not wrong_topics:
        print("[WRONG] No wrong topics found anywhere.")
        return """<div class="section">
          <div class="section-title">🔁 Wrong Answer Review</div>
          <p>No wrong answers found. Complete a quiz first, then click
          <strong>Save Wrong Answers for FRIDAY</strong> when it appears,
          then say <em>explain my wrong answers</em>.</p>
        </div>"""

    # update session
    SESSION["last_wrong_topics"] = wrong_topics

    quiz_topic = json_data.get("topic", SESSION.get("last_topic", "the quiz")) if json_data else SESSION.get("last_topic", "the quiz")
    score      = json_data.get("score", "?") if json_data else "?"
    total      = json_data.get("total", "?") if json_data else "?"

    html = f'''<div class="section">
      <div class="section-title">🔁 Wrong Answer Review</div>
      <p>Quiz topic: <strong>{quiz_topic}</strong> | Score: <strong>{score} / {total}</strong></p>
      <p>Reviewing <strong>{len(wrong_topics)}</strong> topic(s) you got wrong:</p>
    '''

    for topic in wrong_topics:
        print(f"[WRONG] Explaining: '{topic}'")
        prompt = f"""Explain for a Grade {GRADE} CBSE student who got this wrong in a quiz:
"{topic}"

Write 3-4 clear paragraphs explaining:
1. What the correct concept is
2. Why students commonly get it wrong
3. The key rule or formula to remember
4. A quick example

Use LaTeX for all equations: inline $...$ and display $$...$$."""
        explanation = ask_ai(prompt) or "Could not generate explanation."
        html += f"""
        <div class="wrong-review-card">
          <div class="wrong-topic-label">Topic: {topic}</div>
          <div class="wrong-explanation">{explanation.replace(chr(10), "<br>")}</div>
        </div>"""

    html += "</div>"
    return html

# ================= 🔗 CHAIN EXECUTOR =================

def execute_chain(command, chains, user_screen_prompt=None):
    print(f"\n[CHAIN EXEC] Chains: {chains} | Command: '{command}'")

    subject    = detect_subject(command)
    page_title = "FRIDAY Study"
    full_body  = ""
    spoken_bits = []

    # clean topic from command
    topic = command
    for w in ["solve", "explain", "quiz me on", "simulate", "plot", "graph of",
              "on my screen", "on screen", "my screen", "from my screen", "scan",
              "and quiz me", "and explain", "and simulate", "and graph",
              "and plot", "step by step", "walk me through", "similar problems",
              "practice problems", "explain my wrong answers", "the question",
              "read", "check", "look at", "and"]:
        topic = topic.replace(w, "").strip()
    topic = re.sub(r"\s+", " ", topic).strip(" ,.")

    # if screen is involved, topic will be set from user prompt after screenshot
    # for now use last session topic as placeholder
    if not topic or len(topic) < 3:
        topic = SESSION.get("last_topic") or "the question on screen"

    print(f"[CHAIN EXEC] Cleaned topic: '{topic}'")

    # ── screen capture ──
    if "screen_capture" in chains:
        image_path, user_screen_prompt = take_screenshot_with_prompt()
        if user_screen_prompt and len(user_screen_prompt) > 3:
            topic   = user_screen_prompt
            subject = detect_subject(user_screen_prompt) or subject
            print(f"[CHAIN EXEC] Topic updated from user prompt: '{topic}'")

            # re-detect chains from user prompt and merge new ones in
            extra_chains = detect_chains(user_screen_prompt) or []
            for c in extra_chains:
                if c not in chains and c != "screen_capture":
                    chains.append(c)
                    print(f"[CHAIN EXEC] Added chain from user prompt: '{c}'")

        spoken_bits.append("screen captured")

    # ── solve ──
    if "solve" in chains:
        speak("Solving now.")
        from_screen = "screen_capture" in chains
        sol_html, detected_topic = _build_solution(
            "__screen__" if from_screen else topic,
            from_screen=from_screen,
            user_screen_prompt=user_screen_prompt
        )
        full_body  += sol_html
        topic       = detected_topic or topic
        page_title  = "Solution: " + topic[:40]
        spoken_bits.append("solution ready")

        if not SESSION["last_solution"]:
            speak("I could not read the screen. Please try again.")
            return False

    # ── explain ──
    if "explain" in chains:
        speak("Building explanation.")
        full_body  += _build_explanation(topic, subject)
        page_title  = page_title or "Explanation: " + topic[:40]
        spoken_bits.append("explanation added")

    # ── quiz ──
    if "quiz" in chains:
        speak("Generating quiz on " + topic + ".")
        num_q, dur_m = parse_quiz_params(command)
        quiz_html, _ = _build_quiz(topic, subject, num_q, dur_m)
        full_body   += quiz_html
        spoken_bits.append("quiz ready")

    # ── simulate ──
    if "simulate" in chains:
        speak("Building simulation.")
        sim_topic = SESSION.get("last_topic") or topic
        print(f"[CHAIN EXEC] Simulation topic: '{sim_topic}'")
        sim_html = _build_simulation(sim_topic)
        if sim_html and "Could not generate" not in sim_html:
            full_body += sim_html
            spoken_bits.append("simulation added")
        else:
            print("[CHAIN EXEC] Simulation generation failed — skipping.")

    # ── graph ──
    if "graph" in chains:
        speak("Plotting graph.")
        # use detected topic from solve if available, else fall back to session
        graph_topic = SESSION.get("last_topic") or topic
        print(f"[CHAIN EXEC] Graph topic: '{graph_topic}'")
        graph_html = _build_graph(graph_topic)
        if graph_html and "Could not generate" not in graph_html:
            full_body += graph_html
            spoken_bits.append("graph added")
        else:
            print("[CHAIN EXEC] Graph generation failed — skipping.")

    # ── similar problems ──
    if "similar_problems" in chains:
        speak("Generating similar problems.")
        full_body  += _build_similar_problems(topic, subject)
        spoken_bits.append("practice problems added")

    # ── step by step ──
    if "step_by_step" in chains and "solve" not in chains:
        speak("Working through this step by step.")
        from_screen = "screen_capture" in chains
        sol_html, detected_topic = _build_solution(
            "__screen__" if from_screen else topic,
            from_screen=from_screen,
            extra_instruction="Be extremely detailed. Show every single algebraic step.",
            user_screen_prompt=user_screen_prompt
        )
        full_body  += sol_html
        topic       = detected_topic or topic
        spoken_bits.append("step by step solution ready")

    # ── explain wrong answers ──
    if "explain_wrong" in chains:
        speak("Explaining your wrong answers.")
        full_body  += _build_wrong_answer_explanations()
        spoken_bits.append("wrong answer review added")

    if not full_body:
        print("[CHAIN EXEC] No content generated.")
        return False

    html = wrap_html(page_title, full_body)
    open_in_browser(html, page_title[:30].replace(" ", "_").replace(":", ""))

    speak("Done. " + " and ".join(spoken_bits) + ". Everything is open in your browser.")
    print(f"[CHAIN EXEC] Complete: {spoken_bits}")
    return True

# ================= 🎯 SINGLE FEATURE FALLBACKS =================


# ================= 🔍 PDF FINDER =================

def find_pdf_by_description(description):
    """
    Find a PDF/doc by natural language description.
    Searches by keyword + common paths mentioned by user.
    Returns file path or None.
    """
    description = description.lower().strip()

    # ── extract path hints ──
    path_hints = []

    path_keywords = {
        "documents": os.path.expanduser("~/Documents"),
        "downloads": os.path.expanduser("~/Downloads"),
        "desktop":   os.path.expanduser("~/Desktop"),
        "d drive":   "D:\\",
        "d:":        "D:\\",
    }

    search_roots = []
    for keyword, path in path_keywords.items():
        if keyword in description:
            search_roots.append(path)

    # also search the watch folder and common study folders
    search_roots += [
        WATCH_FOLDER,
        os.path.expanduser("~/Documents"),
        os.path.expanduser("~/Downloads"),
        os.path.expanduser("~/Desktop"),
        "D:\\Documents",
        "D:\\Downloads",
        "D:\\Desktop",
    ]

    # deduplicate
    search_roots = list(dict.fromkeys(search_roots))

    # ── extract filename keywords ──
    # remove path/filler words to get subject/filename keywords
    filler = ["the", "pdf", "file", "document", "located", "in", "my",
              "at", "on", "from", "a", "an", "for", "about",
              "documents", "downloads", "desktop", "d drive",
              "math", "maths", "physics", "chemistry", "english", "cs",
              "homework", "assignment", "notes", "worksheet"]

    words = description.split()
    keywords = [w for w in words if w not in filler and len(w) > 2]

    # also add subject keywords back as search terms
    subject_keywords = {
        "math": ["math", "maths", "mathematics"],
        "physics": ["physics", "phy"],
        "chemistry": ["chemistry", "chem"],
        "english": ["english", "eng"],
        "cs": ["cs", "computer"],
    }
    for subj, variants in subject_keywords.items():
        if any(v in description for v in variants):
            keywords.append(subj)

    print(f"[PDF FINDER] Keywords: {keywords}")
    print(f"[PDF FINDER] Search roots: {search_roots}")

    SUPPORTED = [".pdf", ".doc", ".docx", ".txt", ".png", ".jpg", ".jpeg"]
    candidates = []

    for root in search_roots:
        if not os.path.exists(root):
            continue
        for dirpath, dirs, files in os.walk(root):
            # skip hidden and system folders
            dirs[:] = [d for d in dirs if not d.startswith(".") and
                       d.lower() not in ["appdata", "windows", "program files"]]
            for fname in files:
                ext = os.path.splitext(fname)[1].lower()
                if ext not in SUPPORTED:
                    continue
                fname_lower = fname.lower()
                # score by how many keywords match filename
                score = sum(1 for kw in keywords if kw in fname_lower)
                if score > 0:
                    full_path = os.path.join(dirpath, fname)
                    mtime = os.path.getmtime(full_path)
                    candidates.append((score, mtime, full_path))

    if not candidates:
        print("[PDF FINDER] No matching files found.")
        return None

    # sort by score desc, then by most recent
    candidates.sort(key=lambda x: (x[0], x[1]), reverse=True)

    print(f"[PDF FINDER] Top candidates:")
    for score, mtime, path in candidates[:5]:
        print(f"  score={score} | {path}")

    return candidates[0][2]

# ================= 📄 PDF HOMEWORK SOLVER =================

def solve_pdf_homework(file_path):
    """
    Read a PDF/doc, extract all questions, solve each one,
    and display full solutions in browser.
    """
    print(f"[HW SOLVER] Solving: {file_path}")
    speak("Reading your homework file now.")

    # ── read content ──
    ext = os.path.splitext(file_path)[1].lower()
    content = ""

    if ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(file_path)
            for page in doc:
                content += page.get_text()
            doc.close()
            print(f"[HW SOLVER] PDF text: {len(content)} chars")
        except ImportError:
            # fallback to vision on first page
            speak("Reading PDF visually.")
            content = ask_ai_with_image(
                "Read and transcribe all text from this document completely.",
                file_path
            )
    elif ext in [".doc", ".docx"]:
        try:
            import docx
            doc     = docx.Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            print(f"[HW SOLVER] DOCX error: {e}")
    elif ext in [".png", ".jpg", ".jpeg"]:
        content = ask_ai_with_image(
            "Read and transcribe all text from this image completely.",
            file_path
        )
    elif ext == ".txt":
        with open(file_path, "r", errors="ignore") as f:
            content = f.read()

    if not content or len(content) < 20:
        speak("I could not read the file content. Make sure it is not a scanned image PDF.")
        return False

    # ── extract questions from entire file in chunks ──
    speak("Extracting questions from the file.")

    CHUNK_SIZE  = 5000
    chunks      = [content[i:i+CHUNK_SIZE] for i in range(0, len(content), CHUNK_SIZE)]
    all_questions = []

    print(f"[HW SOLVER] File split into {len(chunks)} chunks.")

    for ci, chunk in enumerate(chunks):
        print(f"[HW SOLVER] Extracting from chunk {ci+1}/{len(chunks)}...")

        extract_prompt = f"""You are reading part {ci+1} of {len(chunks)} of a Grade 11 CBSE homework file.

Content:
{chunk}

Extract ALL questions, problems, and exercises from this chunk.
Include numbered questions, lettered sub-parts, fill-in-the-blanks, and word problems.
Skip headings, instructions, and non-question text.

Return ONLY a JSON array of question strings:
["question text", ...]

If no questions in this chunk, return empty array: []
Return ONLY valid JSON array, no markdown."""

        raw = ask_ai(extract_prompt)

        try:
            clean  = re.sub(r"```json|```", "", raw).strip()
            chunk_questions = json.loads(clean)
            if isinstance(chunk_questions, list):
                all_questions.extend(chunk_questions)
                print(f"[HW SOLVER] Chunk {ci+1}: found {len(chunk_questions)} questions.")
        except Exception as e:
            print(f"[HW SOLVER] Chunk {ci+1} parse failed: {e}")

    # deduplicate questions (same question may appear across chunk boundaries)
    seen = set()
    questions = []
    for q in all_questions:
        q_clean = q.strip().lower()[:80]
        if q_clean and q_clean not in seen:
            seen.add(q_clean)
            questions.append(q)

    print(f"[HW SOLVER] Total unique questions: {len(questions)}")

    if not questions:
        speak("I could not find any questions in this file.")
        return False

    speak(f"Found {len(questions)} question{'s' if len(questions) > 1 else ''}. Solving now.")

    # ── solve each question ──
    solutions_html = ""

    for i, question in enumerate(questions):
        print(f"[HW SOLVER] Solving Q{i+1}: {question[:80]}...")

        solve_prompt = f"""You are a Grade 11 CBSE tutor solving a homework question.

Question {i+1}: {question}

Provide:
1. Identify the topic and CBSE chapter
2. Full step-by-step solution with all working shown
3. Final answer clearly stated
4. Key formula or concept used

Use LaTeX for equations: inline $...$ and display $$...$$
Be thorough but clear."""

        solution = ask_ai(solve_prompt, max_tokens=1000)

        if not solution:
            solution = "Could not solve this question. Please try again."

        solutions_html += f"""
        <div class="section">
          <div class="section-title">Question {i+1}</div>
          <div style="background:#1a2e1a; border-left:3px solid #7eff7e;
               padding:12px 16px; border-radius:4px; margin-bottom:15px;
               color:#c0ffc0; font-size:14px;">
            {question}
          </div>
          <div class="step">{clean_latex(solution).replace(chr(10), "<br>")}</div>
        </div>"""

    # ── build page ──
    fname   = os.path.basename(file_path)
    subject = detect_subject(content[:500])

    body = f"""
    <p style="color:#555577; font-size:12px; margin-bottom:20px;">
      File: {fname} | Subject: {subject.title()} | Questions: {len(questions)}
    </p>
    {solutions_html}"""

    html = wrap_html(f"Homework Solutions — {fname}", body)
    open_in_browser(html, "Homework_Solutions")

    speak(f"Done. All {len(questions)} solutions are open in your browser.")
    SESSION["last_topic"]   = subject
    SESSION["last_subject"] = subject
    return True

# ================= 📂 RECENT FILE FINDER =================

def find_recent_file(hint: str = "") -> str | None:
    """
    Find the most recently modified supported file across common folders.
    If a hint is given (e.g. 'physics', 'chemistry'), score by keyword match.
    Supported: .pdf, .doc, .docx, .png, .jpg, .jpeg, .webp
    Returns the best matching file path, or None.
    """
    SUPPORTED_EXTS = {".pdf", ".doc", ".docx", ".png", ".jpg", ".jpeg", ".webp"}

    SEARCH_ROOTS = []
    # Drive D (common on your machine per memory.json)
    for sub in ["Downloads", "Documents", "Desktop"]:
        SEARCH_ROOTS.append(f"D:\\{sub}")
    # Drive C / user home
    SEARCH_ROOTS += [
        os.path.expanduser("D:/Downloads"),
        os.path.expanduser("D:/Documents"),
        os.path.expanduser("D:/Desktop"),
    ]

    # deduplicate while preserving order
    seen_roots = set()
    unique_roots = []
    for r in SEARCH_ROOTS:
        if r not in seen_roots:
            seen_roots.add(r)
            unique_roots.append(r)

    # build hint keywords (strip common noise words)
    NOISE = {"the", "a", "an", "my", "this", "that", "pdf", "file", "document",
             "picture", "image", "photo", "word", "doc", "recent", "latest",
             "newest", "last", "solve", "open", "read"}
    hint_words = [w for w in hint.lower().split() if w not in NOISE and len(w) > 1]

    print(f"[RECENT FILE] Hint keywords: {hint_words}")
    print(f"[RECENT FILE] Searching: {unique_roots}")

    candidates = []   # (keyword_score, mtime, path)

    for root in unique_roots:
        if not os.path.isdir(root):
            continue
        for dirpath, dirs, files in os.walk(root):
            # skip hidden / system dirs
            dirs[:] = [d for d in dirs
                       if not d.startswith(".")
                       and d.lower() not in {"appdata", "windows", "program files",
                                             "system32", "__pycache__"}]
            for fname in files:
                ext = os.path.splitext(fname)[1].lower()
                if ext not in SUPPORTED_EXTS:
                    continue
                full_path   = os.path.join(dirpath, fname)
                mtime       = os.path.getmtime(full_path)
                fname_lower = fname.lower()
                kw_score    = sum(1 for kw in hint_words if kw in fname_lower)
                candidates.append((kw_score, mtime, full_path))

    if not candidates:
        print("[RECENT FILE] No supported files found.")
        return None

    # primary sort: keyword score desc → mtime desc (most recent first)
    candidates.sort(key=lambda x: (x[0], x[1]), reverse=True)

    print("[RECENT FILE] Top 5 candidates:")
    for score, mtime, path in candidates[:5]:
        print(f"  score={score} | mtime={datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M')} | {path}")

    return candidates[0][2]


# ================= 📄 RECENT FILE SOLVER =================

def _extract_file_content(file_path: str) -> tuple[str, list[str]]:
    """
    Extract text/content from a file.
    Returns (text_content, list_of_image_paths).
    For images: text_content is empty, image_paths contains the file.
    For PDFs with no text layer: falls back to page images.
    For DOCX: extracts paragraph + table text.
    """
    ext          = os.path.splitext(file_path)[1].lower()
    text_content = ""
    image_paths  = []

    # ── IMAGE ──
    if ext in {".png", ".jpg", ".jpeg", ".webp"}:
        print(f"[EXTRACT] Image file: {file_path}")
        image_paths.append(file_path)
        return text_content, image_paths

    # ── PDF ──
    if ext == ".pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)

            # try text extraction first
            all_text = ""
            for page in doc:
                all_text += page.get_text()

            if len(all_text.strip()) > 50:
                text_content = all_text
                print(f"[EXTRACT] PDF text extracted: {len(text_content)} chars")
            else:
                # scanned PDF — render each page as image
                print("[EXTRACT] PDF appears scanned — rendering pages as images.")
                speak("This looks like a scanned PDF. Reading it visually.")
                for i, page in enumerate(doc):
                    pix       = page.get_pixmap(dpi=150)
                    img_path  = os.path.join(OUTPUT_DIR, f"pdf_page_{i+1}.png")
                    pix.save(img_path)
                    image_paths.append(img_path)
                    if i >= 4:   # cap at 5 pages to stay within vision limits
                        print("[EXTRACT] Capped at 5 pages for vision.")
                        break

            doc.close()
        except ImportError:
            # fallback: pypdf
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text_content += page.extract_text() or ""
                print(f"[EXTRACT] PDF text (pypdf): {len(text_content)} chars")
            except Exception as e:
                print(f"[EXTRACT] PDF read failed: {e}")
        except Exception as e:
            print(f"[EXTRACT] fitz error: {e}")
        return text_content, image_paths

    # ── DOCX / DOC ──
    if ext in {".docx", ".doc"}:
        try:
            import docx as _docx
            doc   = _docx.Document(file_path)
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            # also grab table cell text
            for table in doc.tables:
                for row in table.rows:
                    lines.append(" | ".join(c.text for c in row.cells if c.text.strip()))
            text_content = "\n".join(lines)
            print(f"[EXTRACT] DOCX text: {len(text_content)} chars")
        except Exception as e:
            print(f"[EXTRACT] DOCX read failed: {e}")
        return text_content, image_paths

    # ── TXT ──
    if ext == ".txt":
        try:
            with open(file_path, "r", errors="ignore") as f:
                text_content = f.read()
        except Exception as e:
            print(f"[EXTRACT] TXT read failed: {e}")
        return text_content, image_paths

    print(f"[EXTRACT] Unsupported extension: {ext}")
    return text_content, image_paths


def _solve_content_with_groq(text_content: str, image_paths: list[str],
                              file_name: str) -> str:
    """
    Send all content to Groq and return a rich, full solution as HTML body string.
    Uses vision model for images/scanned pages, text model for extracted text.
    Handles multi-page documents by chunking.
    """
    subject = detect_subject(text_content[:500] if text_content else file_name)

    SYSTEM_PROMPT = (
        f"You are an expert {BOARD} Grade {GRADE} tutor covering Maths, Physics, Chemistry, "
        f"English, and Computer Science. "
        f"A student has shared a file with you. Your job is to:\n"
        f"1. Identify every question, problem, or exercise in the content.\n"
        f"2. Solve EACH one completely with full step-by-step working.\n"
        f"3. State each final answer clearly.\n"
        f"4. Label each question with its topic/CBSE chapter.\n"
        f"Use LaTeX for math: inline $...$ and display $$...$$.\n"
        f"Format: Use 'Question N:' as a heading for each question, "
        f"then 'Solution:' with the full working, then 'Answer:' in bold."
    )

    parts_html = ""

    # ── TEXT-BASED (PDF text / DOCX / TXT) ──
    if text_content and len(text_content.strip()) > 30:
        CHUNK_SIZE   = 5000
        chunks       = [text_content[i:i+CHUNK_SIZE]
                        for i in range(0, len(text_content), CHUNK_SIZE)]
        total_chunks = len(chunks)
        print(f"[SOLVE FILE] {total_chunks} text chunk(s) to process.")

        if total_chunks > 1:
            speak(f"The file has a lot of content. Solving in {total_chunks} parts.")

        all_solutions = []
        for ci, chunk in enumerate(chunks):
            if total_chunks > 1:
                print(f"[SOLVE FILE] Processing chunk {ci+1}/{total_chunks}...")
            prompt = (
                f"{SYSTEM_PROMPT}\n\n"
                f"--- FILE CONTENT (part {ci+1} of {total_chunks}) ---\n"
                f"{chunk}\n\n"
                f"Solve all questions found in this section."
            )
            result = ask_ai(prompt)
            if result:
                all_solutions.append(result)

        combined = "\n\n---\n\n".join(all_solutions) if all_solutions else ""

        if combined:
            parts_html += f"""
            <div class="section">
              <div class="section-title">📄 Solutions — {os.path.basename(file_name)}</div>
              <div class="step">{clean_latex(combined).replace(chr(10), "<br>")}</div>
            </div>"""

    # ── IMAGE-BASED (scanned PDF pages / photo / screenshot) ──
    if image_paths:
        total_images = len(image_paths)
        speak(f"Reading {'the image' if total_images == 1 else str(total_images) + ' pages'} visually.")
        print(f"[SOLVE FILE] Sending {total_images} image(s) to vision model.")

        for ii, img_path in enumerate(image_paths):
            page_label = f"Page {ii+1}" if total_images > 1 else "Image"
            print(f"[SOLVE FILE] Vision call {ii+1}/{total_images}: {img_path}")

            vision_prompt = (
                f"{SYSTEM_PROMPT}\n\n"
                f"You are looking at {page_label} of the student's file.\n"
                f"Transcribe any questions visible, then solve each one completely."
            )
            result = ask_ai_with_image(vision_prompt, img_path)

            if result:
                parts_html += f"""
                <div class="section">
                  <div class="section-title">🖼️ {page_label} — Solutions</div>
                  <div class="step">{clean_latex(result).replace(chr(10), "<br>")}</div>
                </div>"""

    return parts_html


def solve_recent_file(command: str) -> bool:
    """
    Main entry point for 'solve my recent PDF/document/image' commands.
    Finds the most recent supported file, extracts its content,
    sends it to Groq, and opens the full solution in the browser as HTML.
    """
    print(f"[SOLVE RECENT] Command: '{command}'")

    # ── determine if user specified a file type ──
    cmd = command.lower()

    # extract any hint words (subject, filename keywords)
    hint = cmd
    for noise in ["solve", "open", "read", "my", "the", "recent", "latest",
                  "newest", "last", "pdf", "document", "word", "doc",
                  "picture", "image", "photo", "file", "worksheet",
                  "and solve", "please", "for me", "friday"]:
        hint = hint.replace(noise, " ")
    hint = re.sub(r"\s+", " ", hint).strip()

    # ── decide which extensions to prioritise ──
    wants_image = any(w in cmd for w in ["picture", "image", "photo", "jpeg",
                                          "jpg", "png", "screenshot"])
    wants_doc   = any(w in cmd for w in ["word", "docx", "doc", "document"])
    wants_pdf   = any(w in cmd for w in ["pdf"])

    speak("Looking for your most recent file.")
    file_path = find_recent_file(hint)

    if not file_path:
        speak("I could not find any recent files in your Downloads, Documents, or Desktop. "
              "Make sure the file is there and try again.")
        return False

    ext       = os.path.splitext(file_path)[1].lower()
    fname     = os.path.basename(file_path)
    type_name = {".pdf": "PDF", ".doc": "Word document", ".docx": "Word document",
                 ".png": "image", ".jpg": "image", ".jpeg": "image",
                 ".webp": "image"}.get(ext, "file")

    speak(f"Found {type_name}: {fname}. Solving it now. This might take a moment.")
    print(f"[SOLVE RECENT] Using: {file_path}")

    # ── extract content ──
    text_content, image_paths = _extract_file_content(file_path)

    if not text_content and not image_paths:
        speak("I could not read that file. It may be corrupted or an unsupported format.")
        return False

    # ── solve with Groq ──
    speak("Sending to Groq. Solving all questions now.")
    solutions_html = _solve_content_with_groq(text_content, image_paths, file_path)

    if not solutions_html:
        speak("Groq could not find any questions to solve in this file.")
        return False

    # ── build metadata banner ──
    mtime_str = datetime.fromtimestamp(os.path.getmtime(file_path)).strftime("%d %b %Y, %H:%M")
    subject   = detect_subject(text_content[:300] if text_content else fname)
    meta_html = f"""
    <div style="background:#13131f; border:1px solid #2a2a4e; border-radius:8px;
                padding:14px 20px; margin-bottom:24px; font-size:13px; color:#7eb8ff;">
      <strong>📁 File:</strong> {fname} &nbsp;|&nbsp;
      <strong>Type:</strong> {type_name} &nbsp;|&nbsp;
      <strong>Modified:</strong> {mtime_str} &nbsp;|&nbsp;
      <strong>Subject:</strong> {subject.title()}
    </div>"""

    full_body  = meta_html + solutions_html
    page_title = f"Solutions — {fname}"
    html       = wrap_html(page_title, full_body)
    open_in_browser(html, "Recent_File_Solution")

    SESSION["last_topic"]   = subject
    SESSION["last_subject"] = subject

    speak(f"Done. All solutions for {fname} are open in your browser. "
          f"Say quiz me on this if you want to practice.")
    return True


def solve_homework(command):
    """
    Ask user if homework is a PDF or on screen.
    If PDF, ask for file description and search for it.
    If screen, take screenshot and solve.
    """
    speak("Is your homework a PDF or document, or is it on your screen?")
    response = listen_for_prompt()

    if not response:
        speak("I did not catch that. I will scan your screen.")
        response = "screen"

    # ── PDF / file path ──
    if any(w in response for w in ["pdf", "file", "document", "doc", "folder",
                                    "documents", "downloads", "desktop", "drive",
                                    "word", "located", "in my", "saved"]):
        speak("What is the file name or where is it located? For example: the math PDF in my documents.")
        file_description = listen_for_prompt(timeout=12)

        if not file_description:
            speak("I did not catch that. Please try again.")
            return True

        print(f"[HW SOLVER] Looking for: {file_description}")
        speak("Searching for the file.")

        file_path = find_pdf_by_description(file_description)

        if not file_path:
            speak("I could not find the file. Make sure it is in Documents, Downloads, or Desktop and try again with a clearer description.")
            return True

        fname = os.path.basename(file_path)
        speak(f"Found it: {fname}. Solving all questions now.")
        return solve_pdf_homework(file_path)

    # ── screen ──
    else:
        image_path, user_prompt = take_screenshot_with_prompt()
        speak("Analyzing now.")
        sol_html, topic = _build_solution("__screen__", from_screen=True,
                                          user_screen_prompt=user_prompt)
        if not SESSION["last_solution"]:
            speak("I could not read the screen. Please try again.")
            return False
        html = wrap_html("Homework Solution", sol_html)
        open_in_browser(html, "Homework_Solution")
        speak("Solution is open in your browser. Say quiz me on this if you want to practice.")
        return True

def explain_concept(command, subject=None):
    topic = command
    for w in ["explain", "what is", "tell me about", "teach me", "how does", "how do"]:
        topic = topic.replace(w, "").strip()
    speak("Preparing explanation for " + topic + ".")
    exp_html = _build_explanation(topic, subject)
    html     = wrap_html("Explanation: " + topic, exp_html)
    open_in_browser(html, "Explain_" + topic[:20].replace(" ", "_"))
    speak("Explanation is ready. Want me to quiz you on this?")
    return True

def solve_equation(command):
    eq = command
    for w in ["solve", "calculate", "find", "evaluate", "differentiate",
              "integrate", "simplify", "factorise", "expand", "prove"]:
        eq = eq.replace(w, "").strip()
    speak("Solving.")
    sol_html, _ = _build_solution(eq)
    html = wrap_html("Solution", sol_html)
    open_in_browser(html, "Equation")
    speak("Solution is open in your browser.")
    return True

def generate_graph(command):
    speak("Generating graph.")
    graph_html = _build_graph(command)
    html       = wrap_html("Graph", graph_html)
    open_in_browser(html, "Graph")
    speak("Graph is open in your browser.")
    return True

def run_simulation(command):
    speak("Building simulation.")
    sim_html = _build_simulation(command)
    html     = wrap_html("Simulation", sim_html)
    open_in_browser(html, "Simulation")
    speak("Simulation is running in your browser.")
    return True

def generate_quiz(command, subject=None):
    num_questions, duration_minutes = parse_quiz_params(command)
    topic = command
    for w in ["quiz", "test", "questions", "me on", "about", "give me", "ask me",
              "for", "minutes", "minute", "mins", "min", "mcqs", "mcq"]:
        topic = topic.replace(w, "").strip()
    # strip numbers (question count / duration)
    topic = re.sub(r"\b\d+\b", "", topic).strip(" ,.")
    if not topic or len(topic) < 3:
        topic = SESSION.get("last_topic") or "general science"
    speak("Preparing quiz on " + topic + ".")
    quiz_html, _ = _build_quiz(topic, subject, num_questions, duration_minutes)
    timer_label  = f" ({duration_minutes} min)" if duration_minutes else ""
    html         = wrap_html("Quiz: " + topic + timer_label, quiz_html)
    open_in_browser(html, "Quiz_" + topic[:15].replace(" ", "_"))
    speak("Quiz is ready. Good luck!")
    return True

def solve_puzzle(command):
    puzzle = command
    for w in ["solve", "crack", "figure out", "puzzle"]:
        puzzle = puzzle.replace(w, "").strip()
    speak("Working on this puzzle.")
    sol_html, _ = _build_solution(puzzle)
    html        = wrap_html("Puzzle Solution", sol_html)
    open_in_browser(html, "Puzzle")
    speak("Solution is open in your browser.")
    return True

def explain_wrong_answers():
    speak("Pulling up your wrong answers from the last quiz.")
    wrong_html = _build_wrong_answer_explanations()
    html       = wrap_html("Wrong Answer Review", wrong_html)
    open_in_browser(html, "Wrong_Answer_Review")
    speak("Review is open in your browser.")
    return True

# ================= 🎯 COMMAND ROUTING =================

STUDY_TRIGGERS = {
    # homework MUST come before equation so "solve my homework" doesnt match "solve"
    "homework":   ["solve my homework", "help with homework", "scan my homework",
                   "do my homework", "my homework", "homework help",
                   "answer my homework", "finish my homework"],
    "explain":    ["explain", "what is", "tell me about", "teach me",
                   "how does", "how do", "what are", "define"],
    "equation":   ["solve", "calculate", "find the value", "evaluate",
                   "differentiate", "integrate", "simplify", "factorise",
                   "expand", "prove that"],
    "graph":      ["plot", "draw a graph", "graph of", "show me the graph",
                   "make a graph", "visualize", "chart"],
    "simulate":   ["simulate", "show me a simulation", "run a simulation",
                   "animate", "demonstrate"],
    "quiz":       ["quiz me", "test me", "give me questions", "ask me questions",
                   "practice questions", "mcq", "multiple choice"],
    "puzzle":     ["solve this puzzle", "crack this", "brain teaser", "logic puzzle"],
    "wrong":      ["explain wrong", "wrong answers", "what i got wrong", "my mistakes"],
}

def route(command):
    """NLP-based routing within study mode."""
    intent, score, all_scores = classify(command, "study", threshold=0.30)

    print(f"[STUDY ROUTE] NLP intent={intent} score={score:.3f}")

    if intent:
        return intent

    # fallback to keyword matching for very short commands
    cmd = command.lower()
    for action, triggers in STUDY_TRIGGERS.items():
        for trigger in triggers:
            if trigger in cmd:
                print(f"[STUDY ROUTE] Keyword fallback: '{action}' via '{trigger}'")
                return action

    print("[STUDY ROUTE] No match — defaulting to explain")
    return "explain"

# ================= 🎯 MAIN HANDLE =================

def handle(command):

    command = command.lower().strip()

    for prefix in ["study mode", "friday"]:
        command = command.replace(prefix, "").strip()

    print(f"\n{'='*60}")
    print(f"[HANDLE] Command: '{command}'")
    print(f"{'='*60}")

    chains = detect_chains(command)

    if chains and len(chains) > 1:
        print(f"[HANDLE] Multi-chain: {chains}")
        return execute_chain(command, chains)

    action  = route(command)
    subject = detect_subject(command)

    print(f"[HANDLE] Single action: '{action}' | subject: '{subject}'")

    if action == "homework":
        return solve_homework(command)

    if action == "explain":
        return explain_concept(command, subject)

    if action == "equation":
        if any(w in command for w in ["screen", "this", "it"]):
            return execute_chain(command, ["screen_capture", "solve"])
        return solve_equation(command)

    if action == "graph":
        return generate_graph(command)

    if action == "simulate":
        return run_simulation(command)

    if action == "quiz":
        return generate_quiz(command, subject)

    if action == "puzzle":
        return solve_puzzle(command)

    if action == "wrong":
        return explain_wrong_answers()

    return explain_concept(command, subject)